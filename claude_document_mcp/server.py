"""
Claude Document MCP Server - Model Context Protocol server for Claude Desktop

Features:
- Microsoft Word file operations (create, edit, convert from txt)
- Excel file operations (create, edit, convert from csv)
- PDF file operations (create, convert from Word)
- File management with trash support (delete to recycle bin or permanent)

This is a headless server with no UI, designed to be used with Claude Desktop.
"""

import os
import sys
import json
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional

from mcp.server.fastmcp import FastMCP

# Document processing libraries
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise ImportError("Please install python-docx with: pip install python-docx")

try:
    import pandas as pd
    import openpyxl
except ImportError:
    raise ImportError("Please install pandas and openpyxl with: pip install pandas openpyxl")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    raise ImportError("Please install reportlab with: pip install reportlab")

try:
    import docx2pdf
except ImportError:
    raise ImportError("Please install docx2pdf with: pip install docx2pdf")

try:
    from send2trash import send2trash
except ImportError:
    raise ImportError("Please install send2trash with: pip install send2trash")

# Set up logging
log_dir = Path(__file__).parent.parent / "logs"
log_dir.mkdir(exist_ok=True)
log_file = log_dir / "document_mcp.log"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler(log_file)
    ]
)
logger = logging.getLogger(__name__)

# Initialize the FastMCP server
server = FastMCP("Document Operations")
mcp = server

# ---- Helper Functions ----

def _set_cell_shading(cell, color_hex: str):
    """Helper function to set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _fix_cell_paragraph_spacing(cell):
    """Fix paragraph spacing in cell to match Word default (no extra spacing)."""
    for para in cell.paragraphs:
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.0


def _set_cell_borders(cell, color: str = "CCCCCC", size: str = "4"):
    """Helper function to set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_cell_width(cell, width_dxa: int):
    """Helper function to set cell width in DXA (twentieths of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(width_dxa))
    tcPr.append(tcW)


def _set_row_height(row, height_pt: int = None, auto_fit: bool = True):
    """Helper function to set row height."""
    if height_pt is None and auto_fit:
        return
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    if auto_fit:
        trHeight.set(qn('w:val'), str(int((height_pt or 0) * 20)))
        trHeight.set(qn('w:hRule'), 'auto')
    else:
        trHeight.set(qn('w:val'), str(int(height_pt * 20)))
        trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def _set_table_grid(table, col_widths: list):
    """Set table grid column widths and remove cell widths to let Word use grid."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is not None:
            tblW.set(qn('w:w'), '100')
            tblW.set(qn('w:type'), 'auto')
        tblLook = tblPr.find(qn('w:tblLook'))
        if tblLook is not None:
            tblPr.remove(tblLook)
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
    
    for child in list(tbl):
        if child.tag.endswith('tblGrid'):
            tbl.remove(child)
    
    tblGrid = OxmlElement('w:tblGrid')
    for width in col_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width))
        tblGrid.append(gridCol)
    
    if tblPr is not None:
        tblPr_index = list(tbl).index(tblPr)
        tbl.insert(tblPr_index + 1, tblGrid)
    else:
        tbl.insert(0, tblGrid)
    
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)


def _parse_rich_text(paragraph, text: str, default_font: str = 'Arial', default_size: int = 11, default_color: str = None, default_italic: bool = False):
    """Parse text with **bold** markers and add runs to paragraph."""
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = default_font
            run.font.size = Pt(default_size)
            if default_color:
                r, g, b = int(default_color[0:2], 16), int(default_color[2:4], 16), int(default_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            if default_italic:
                run.italic = True
        elif part:
            run = paragraph.add_run(part)
            run.font.name = default_font
            run.font.size = Pt(default_size)
            if default_color:
                r, g, b = int(default_color[0:2], 16), int(default_color[2:4], 16), int(default_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            if default_italic:
                run.italic = True


# ---- Microsoft Word Operations ----

@server.tool()
def create_word_document(filepath: str, content: str) -> Dict[str, Any]:
    """Create a new Microsoft Word document with the provided content."""
    try:
        doc = Document()
        doc.add_paragraph(content)
        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        doc.save(filepath)
        logger.info(f"Created Word document: {filepath}")
        return {"success": True, "message": "Successfully created Word document", "filepath": filepath}
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        return {"success": False, "message": f"Error creating Word document: {str(e)}", "filepath": None}


@server.tool()
def read_word_document_structure(filepath: str) -> Dict[str, Any]:
    """Read and analyze the structure of a Word document, including table properties."""
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        with zipfile.ZipFile(filepath, 'r') as zf:
            with zf.open('word/document.xml') as f:
                tree = ET.parse(f)
                root = tree.getroot()
            has_header = any('header' in name for name in zf.namelist())
            has_footer = any('footer' in name for name in zf.namelist())
        
        tables_info = []
        tables = root.findall('.//w:tbl', ns)
        
        for i, tbl in enumerate(tables):
            table_info = {
                "table_index": i + 1, "column_widths": [], "column_widths_source": None,
                "row_count": 0, "row_heights": [], "has_explicit_row_heights": False
            }
            
            tblGrid = tbl.find('w:tblGrid', ns)
            if tblGrid is not None:
                gridCols = tblGrid.findall('w:gridCol', ns)
                table_info["column_widths"] = [
                    int(gc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', 0))
                    for gc in gridCols
                ]
                table_info["column_widths_source"] = "tblGrid"
            
            rows = tbl.findall('w:tr', ns)
            table_info["row_count"] = len(rows)
            table_info["cell_widths"] = []
            
            for j, row in enumerate(rows):
                row_info = {"row_index": j, "height": None, "height_rule": None}
                trPr = row.find('w:trPr', ns)
                if trPr is not None:
                    trHeight = trPr.find('w:trHeight', ns)
                    if trHeight is not None:
                        table_info["has_explicit_row_heights"] = True
                        val = trHeight.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        hRule = trHeight.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hRule')
                        row_info["height"] = int(val) if val else None
                        row_info["height_rule"] = hRule
                table_info["row_heights"].append(row_info)
                
                cells = row.findall('w:tc', ns)
                cell_widths_row = []
                for cell in cells:
                    tcPr = cell.find('w:tcPr', ns)
                    if tcPr is not None:
                        tcW = tcPr.find('w:tcW', ns)
                        if tcW is not None:
                            cell_widths_row.append(int(tcW.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', 0)))
                        else:
                            cell_widths_row.append(None)
                    else:
                        cell_widths_row.append(None)
                table_info["cell_widths"].append(cell_widths_row)
            
            tables_info.append(table_info)
        
        paragraphs = root.findall('.//w:p', ns)
        
        return {
            "success": True, "filepath": filepath, "tables_count": len(tables_info),
            "tables": tables_info, "paragraphs_count": len(paragraphs),
            "has_header": has_header, "has_footer": has_footer
        }
        
    except Exception as e:
        logger.error(f"Error analyzing Word document: {str(e)}")
        return {"success": False, "message": f"Error analyzing Word document: {str(e)}", "filepath": filepath}


@server.tool()
def compare_word_documents(filepath1: str, filepath2: str) -> Dict[str, Any]:
    """Compare the structure of two Word documents, focusing on table properties."""
    try:
        doc1 = read_word_document_structure(filepath1)
        doc2 = read_word_document_structure(filepath2)
        
        if not doc1.get("success") or not doc2.get("success"):
            return {"success": False, "message": "Failed to read one or both documents",
                    "doc1_error": doc1.get("message"), "doc2_error": doc2.get("message")}
        
        differences = []
        
        if doc1["tables_count"] != doc2["tables_count"]:
            differences.append({"type": "tables_count", "doc1": doc1["tables_count"], "doc2": doc2["tables_count"]})
        
        min_tables = min(doc1["tables_count"], doc2["tables_count"])
        for i in range(min_tables):
            t1, t2 = doc1["tables"][i], doc2["tables"][i]
            table_diffs = {"table_index": i + 1, "differences": []}
            
            if t1["column_widths"] != t2["column_widths"]:
                table_diffs["differences"].append({
                    "property": "column_widths", "doc1": t1["column_widths"], "doc2": t2["column_widths"],
                    "doc1_source": t1["column_widths_source"], "doc2_source": t2["column_widths_source"]
                })
            if t1["has_explicit_row_heights"] != t2["has_explicit_row_heights"]:
                table_diffs["differences"].append({
                    "property": "has_explicit_row_heights",
                    "doc1": t1["has_explicit_row_heights"], "doc2": t2["has_explicit_row_heights"]
                })
            if t1["row_count"] != t2["row_count"]:
                table_diffs["differences"].append({"property": "row_count", "doc1": t1["row_count"], "doc2": t2["row_count"]})
            
            if table_diffs["differences"]:
                differences.append(table_diffs)
        
        is_identical = len(differences) == 0
        return {
            "success": True, "is_identical": is_identical, "filepath1": filepath1, "filepath2": filepath2,
            "differences": differences, "summary": "Documents are identical in structure" if is_identical else f"Found {len(differences)} difference(s)"
        }
        
    except Exception as e:
        logger.error(f"Error comparing Word documents: {str(e)}")
        return {"success": False, "message": f"Error comparing Word documents: {str(e)}"}


@server.tool()
def create_formatted_word_document(filepath: str, document_data: str) -> Dict[str, Any]:
    """
    Create a formatted Microsoft Word document with rich styling support.
    
    Args:
        filepath: Path where to save the document
        document_data: JSON string containing document structure with "sections" key (NOT "content")
    
    IMPORTANT: Use "sections" not "content" in JSON, otherwise the file will be empty!
    
    Example:
        {
            "title": "Document Title",
            "subtitle": "Subtitle",
            "header": "Header text",
            "footer": "Footer text",
            "sections": [
                {"type": "heading", "level": 1, "text": "Section Title"},
                {"type": "paragraph", "text": "Paragraph content"},
                {"type": "bullet_list", "items": ["Item 1", "Item 2"]},
                {"type": "table", "headers": ["Col1", "Col2"], "rows": [["A", "B"]]}
            ]
        }
    """
    try:
        try:
            data = json.loads(document_data)
        except json.JSONDecodeError as e:
            return {"success": False, "message": f"Invalid JSON format: {str(e)}", "filepath": None}

        doc = Document()

        for doc_section in doc.sections:
            doc_section.top_margin = Cm(1.9)
            doc_section.bottom_margin = Cm(1.9)
            doc_section.left_margin = Cm(1.9)
            doc_section.right_margin = Cm(1.9)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        if data.get("header"):
            doc_section = doc.sections[0]
            header = doc_section.header
            header_para = header.paragraphs[0]
            header_para.text = data["header"]
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in header_para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(102, 102, 102)
                run.font.italic = True

        if data.get("footer"):
            doc_section = doc.sections[0]
            footer = doc_section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = data["footer"]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in footer_para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(102, 102, 102)

        if data.get("title"):
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(data["title"])
            title_run.bold = True
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(24)
            title_run.font.color.rgb = RGBColor(31, 78, 121)

        if data.get("subtitle"):
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(data["subtitle"])
            subtitle_run.italic = True
            subtitle_run.font.name = 'Arial'
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
            subtitle_para.paragraph_format.space_after = Pt(20)

        for section in data.get("sections", []):
            section_type = section.get("type")

            if section_type == "heading":
                level = section.get("level", 1)
                text = section.get("text", "")
                color_hex = section.get("color", "1F4E79")
                heading = doc.add_heading(text, level=level)
                for run in heading.runs:
                    run.font.name = 'Arial'
                    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                heading.paragraph_format.space_before = Pt(15)
                heading.paragraph_format.space_after = Pt(6)

            elif section_type == "paragraph":
                text = section.get("text", "")
                para = doc.add_paragraph()
                alignment = section.get("alignment", "left").lower()
                if alignment == "center": para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == "right": para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == "justify": para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else: para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                font_size = section.get("font_size", 11)
                color_hex = section.get("color", None)
                is_italic = section.get("italic", False)
                is_bold = section.get("bold", False)
                
                if '**' in text:
                    _parse_rich_text(para, text, 'Arial', font_size, color_hex, is_italic)
                else:
                    run = para.add_run(text)
                    run.bold = is_bold
                    run.italic = is_italic
                    run.font.name = 'Arial'
                    run.font.size = Pt(font_size)
                    if color_hex:
                        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                para.paragraph_format.space_after = Pt(section.get("space_after", 6))

            elif section_type == "bullet_list":
                for item in section.get("items", []):
                    para = doc.add_paragraph(item, style='List Bullet')
                    for run in para.runs: run.font.name = 'Arial'

            elif section_type == "numbered_list":
                for item in section.get("items", []):
                    para = doc.add_paragraph(item, style='List Number')
                    for run in para.runs: run.font.name = 'Arial'

            elif section_type == "table":
                headers = section.get("headers", [])
                rows = section.get("rows", [])
                header_bg = section.get("header_bg_color", "1F4E79")
                header_text_color = section.get("header_text_color", "FFFFFF")
                alt_row = section.get("alt_row_color", "F2F2F2")
                row_height = section.get("row_height", None)
                col_widths = section.get("col_widths", None)

                if headers:
                    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                    if col_widths: _set_table_grid(table, col_widths)
                    
                    header_row = table.rows[0]
                    if row_height: _set_row_height(header_row, row_height, auto_fit=True)
                    for i, header_text_val in enumerate(headers):
                        cell = header_row.cells[i]
                        cell.text = str(header_text_val)
                        _set_cell_shading(cell, header_bg)
                        _set_cell_borders(cell)
                        _fix_cell_paragraph_spacing(cell)
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
                                run.font.name = 'Arial'
                                r, g, b = int(header_text_color[0:2], 16), int(header_text_color[2:4], 16), int(header_text_color[4:6], 16)
                                run.font.color.rgb = RGBColor(r, g, b)

                    for row_idx, row_data in enumerate(rows):
                        table_row = table.rows[row_idx + 1]
                        if row_height: _set_row_height(table_row, row_height, auto_fit=True)
                        for col_idx, cell_value in enumerate(row_data):
                            if col_idx < len(table_row.cells):
                                cell = table_row.cells[col_idx]
                                cell.text = str(cell_value)
                                _set_cell_borders(cell)
                                _fix_cell_paragraph_spacing(cell)
                                for para in cell.paragraphs:
                                    for run in para.runs: run.font.name = 'Arial'
                                if row_idx % 2 == 1: _set_cell_shading(cell, alt_row)
                    doc.add_paragraph()

            elif section_type == "key_value_table":
                rows = section.get("rows", [])
                first_col_bg = section.get("first_col_bg_color", "D6E3F0")
                first_cell_bg = section.get("first_cell_bg_color", "1F4E79")
                first_cell_text = section.get("first_cell_text_color", "FFFFFF")
                col_widths = section.get("col_widths", None)
                row_height = section.get("row_height", None)

                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows else 2)
                    if col_widths: _set_table_grid(table, col_widths)
                    
                    for row_idx, row_data in enumerate(rows):
                        table_row = table.rows[row_idx]
                        if row_height: _set_row_height(table_row, row_height, auto_fit=True)
                        for col_idx, cell_value in enumerate(row_data):
                            if col_idx < len(table_row.cells):
                                cell = table_row.cells[col_idx]
                                cell.text = str(cell_value)
                                _set_cell_borders(cell)
                                _fix_cell_paragraph_spacing(cell)
                                
                                if row_idx == 0 and col_idx == 0:
                                    _set_cell_shading(cell, first_cell_bg)
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.bold = True
                                            run.font.name = 'Arial'
                                            r, g, b = int(first_cell_text[0:2], 16), int(first_cell_text[2:4], 16), int(first_cell_text[4:6], 16)
                                            run.font.color.rgb = RGBColor(r, g, b)
                                elif col_idx == 0:
                                    _set_cell_shading(cell, first_col_bg)
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.bold = True
                                            run.font.name = 'Arial'
                                else:
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.font.name = 'Arial'
                                            if row_idx == 1 and col_idx == 1: run.bold = True
                    doc.add_paragraph()

            elif section_type == "page_break":
                doc.add_page_break()

            elif section_type == "spacer":
                doc.add_paragraph()

        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        doc.save(filepath)

        logger.info(f"Created formatted Word document: {filepath}")
        return {"success": True, "message": "Successfully created formatted Word document", "filepath": filepath}
    except Exception as e:
        logger.error(f"Error creating formatted Word document: {str(e)}")
        return {"success": False, "message": f"Error creating formatted Word document: {str(e)}", "filepath": None}


@server.tool()
def edit_word_document(filepath: str, operations: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Edit an existing Microsoft Word document using the specified operations."""
    try:
        if not os.path.exists(filepath):
            return {"success": False, "message": f"File not found: {filepath}", "filepath": None}

        doc = Document(filepath)

        for op in operations:
            op_type = op.get("type")
            if op_type == "add_paragraph":
                doc.add_paragraph(op.get("text", ""))
            elif op_type == "add_heading":
                doc.add_heading(op.get("text", ""), level=op.get("level", 1))
            elif op_type == "edit_paragraph":
                idx = op.get("index", 0)
                if 0 <= idx < len(doc.paragraphs):
                    doc.paragraphs[idx].text = op.get("text", "")
            elif op_type == "delete_paragraph":
                idx = op.get("index", 0)
                if 0 <= idx < len(doc.paragraphs):
                    p = doc.paragraphs[idx]
                    p._element.getparent().remove(p._element)

        doc.save(filepath)
        return {"success": True, "message": "Successfully edited Word document", "filepath": filepath}
    except Exception as e:
        return {"success": False, "message": f"Error editing Word document: {str(e)}", "filepath": None}


@server.tool()
def convert_txt_to_word(source_path: str, target_path: str) -> Dict[str, Any]:
    """Convert a text file to a Microsoft Word document."""
    try:
        if not os.path.exists(source_path):
            return {"success": False, "message": f"Source file not found: {source_path}", "filepath": None}
        with open(source_path, 'r', encoding='utf-8') as file:
            text_content = file.read()
        doc = Document()
        for paragraph in text_content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)
        doc.save(target_path)
        return {"success": True, "message": "Successfully converted text to Word document", "filepath": target_path}
    except Exception as e:
        return {"success": False, "message": f"Error converting text to Word: {str(e)}", "filepath": None}


# ---- Excel Operations ----

@server.tool()
def create_excel_file(filepath: str, content: str) -> Dict[str, Any]:
    """Create a new Excel file with the provided content."""
    try:
        try:
            data = json.loads(content)
        except json.JSONDecodeError:
            data = [line.split(',') for line in content.strip().split('\n')]
        df = pd.DataFrame(data)
        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        df.to_excel(filepath, index=False)
        return {"success": True, "message": "Successfully created Excel file", "filepath": filepath}
    except Exception as e:
        return {"success": False, "message": f"Error creating Excel file: {str(e)}", "filepath": None}


@server.tool()
def edit_excel_file(filepath: str, operations: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Edit an existing Excel file using the specified operations."""
    try:
        if not os.path.exists(filepath):
            return {"success": False, "message": f"File not found: {filepath}", "filepath": None}
        wb = openpyxl.load_workbook(filepath)
        for op in operations:
            op_type = op.get("type")
            sheet_name = op.get("sheet", wb.sheetnames[0])
            if sheet_name not in wb.sheetnames:
                wb.create_sheet(sheet_name)
            sheet = wb[sheet_name]
            if op_type == "update_cell":
                sheet.cell(row=op.get("row", 1), column=op.get("col", 1), value=op.get("value", ""))
            elif op_type == "update_range":
                start_row, start_col = op.get("start_row", 1), op.get("start_col", 1)
                for i, row_values in enumerate(op.get("values", [])):
                    for j, value in enumerate(row_values):
                        sheet.cell(row=start_row + i, column=start_col + j, value=value)
            elif op_type == "delete_row":
                sheet.delete_rows(op.get("row", 1))
            elif op_type == "delete_column":
                sheet.delete_cols(op.get("col", 1))
            elif op_type == "add_sheet":
                if op.get("name", "NewSheet") not in wb.sheetnames:
                    wb.create_sheet(op.get("name", "NewSheet"))
            elif op_type == "delete_sheet":
                if sheet_name in wb.sheetnames and len(wb.sheetnames) > 1:
                    del wb[sheet_name]
        wb.save(filepath)
        return {"success": True, "message": "Successfully edited Excel file", "filepath": filepath}
    except Exception as e:
        return {"success": False, "message": f"Error editing Excel file: {str(e)}", "filepath": None}


@server.tool()
def convert_csv_to_excel(source_path: str, target_path: str) -> Dict[str, Any]:
    """Convert a CSV file to an Excel file."""
    try:
        if not os.path.exists(source_path):
            return {"success": False, "message": f"Source file not found: {source_path}", "filepath": None}
        df = pd.read_csv(source_path)
        os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)
        df.to_excel(target_path, index=False)
        return {"success": True, "message": "Successfully converted CSV to Excel", "filepath": target_path}
    except Exception as e:
        return {"success": False, "message": f"Error converting CSV to Excel: {str(e)}", "filepath": None}


# ---- PDF Operations ----

@server.tool()
def create_pdf_file(filepath: str, content: str) -> Dict[str, Any]:
    """Create a new PDF file with the provided text content."""
    try:
        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        c = canvas.Canvas(filepath, pagesize=letter)
        width, height = letter
        lines = content.split('\n')
        y_position = height - 40
        for line in lines:
            if y_position < 40:
                c.showPage()
                y_position = height - 40
            c.drawString(40, y_position, line)
            y_position -= 15
        c.save()
        return {"success": True, "message": "Successfully created PDF file", "filepath": filepath}
    except Exception as e:
        return {"success": False, "message": f"Error creating PDF file: {str(e)}", "filepath": None}


@server.tool()
def convert_word_to_pdf(source_path: str, target_path: str) -> Dict[str, Any]:
    """Convert a Microsoft Word document to a PDF file."""
    try:
        if not os.path.exists(source_path):
            return {"success": False, "message": f"Source file not found: {source_path}", "filepath": None}
        os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)
        docx2pdf.convert(source_path, target_path)
        return {"success": True, "message": "Successfully converted Word to PDF", "filepath": target_path}
    except Exception as e:
        return {"success": False, "message": f"Error converting Word to PDF: {str(e)}", "filepath": None}


# ---- File Management Operations ----

@server.tool()
def delete_file(filepath: str, confirm: str) -> Dict[str, Any]:
    """
    Delete a file with explicit user confirmation.
    
    Args:
        filepath: Path to the file to delete
        confirm: Must be "CORBEILLE" (recoverable) or "SUPPRESSION DÉFINITIVE" (permanent)
    
    Returns:
        Success status and message
    """
    try:
        if confirm not in ["CORBEILLE", "SUPPRESSION DÉFINITIVE"]:
            return {
                "success": False, 
                "message": "Suppression annulee. Le parametre 'confirm' doit etre: 'CORBEILLE' (recuperable) ou 'SUPPRESSION DÉFINITIVE' (permanent)",
                "filepath": filepath,
                "deleted": False
            }
        
        if not os.path.exists(filepath):
            return {"success": False, "message": f"File not found: {filepath}", "filepath": filepath, "deleted": False}
        
        if os.path.isdir(filepath):
            return {"success": False, "message": f"Cannot delete directory with this function: {filepath}", "filepath": filepath, "deleted": False}
        
        file_size = os.path.getsize(filepath)
        file_name = os.path.basename(filepath)
        
        if confirm == "CORBEILLE":
            send2trash(filepath)
            logger.info(f"Sent to trash: {filepath} ({file_size} bytes)")
            return {
                "success": True, 
                "message": f"Envoye a la corbeille: {file_name} ({file_size} bytes) - RECUPERABLE",
                "filepath": filepath,
                "deleted": True,
                "method": "trash"
            }
        else:  # SUPPRESSION DÉFINITIVE
            os.remove(filepath)
            logger.info(f"Permanently deleted: {filepath} ({file_size} bytes)")
            return {
                "success": True, 
                "message": f"Supprime definitivement: {file_name} ({file_size} bytes) - NON RECUPERABLE",
                "filepath": filepath,
                "deleted": True,
                "method": "permanent"
            }
    except PermissionError:
        return {"success": False, "message": f"Permission denied: {filepath}", "filepath": filepath, "deleted": False}
    except Exception as e:
        logger.error(f"Error deleting file: {str(e)}")
        return {"success": False, "message": f"Error deleting file: {str(e)}", "filepath": filepath, "deleted": False}


@server.tool()
def delete_directory(dirpath: str, confirm: str) -> Dict[str, Any]:
    """
    Delete an empty directory with explicit user confirmation.
    
    Args:
        dirpath: Path to the directory to delete (must be empty)
        confirm: Must be "CORBEILLE" (recoverable) or "SUPPRESSION DÉFINITIVE" (permanent)
    
    Returns:
        Success status and message
    """
    try:
        if confirm not in ["CORBEILLE", "SUPPRESSION DÉFINITIVE"]:
            return {
                "success": False, 
                "message": "Suppression annulee. Le parametre 'confirm' doit etre: 'CORBEILLE' (recuperable) ou 'SUPPRESSION DÉFINITIVE' (permanent)",
                "dirpath": dirpath,
                "deleted": False
            }
        
        if not os.path.exists(dirpath):
            return {"success": False, "message": f"Directory not found: {dirpath}", "dirpath": dirpath, "deleted": False}
        
        if not os.path.isdir(dirpath):
            return {"success": False, "message": f"Not a directory: {dirpath}", "dirpath": dirpath, "deleted": False}
        
        contents = os.listdir(dirpath)
        if contents:
            return {
                "success": False, 
                "message": f"Directory not empty ({len(contents)} items). Delete contents first.",
                "dirpath": dirpath,
                "contents": contents[:10],
                "deleted": False
            }
        
        if confirm == "CORBEILLE":
            send2trash(dirpath)
            logger.info(f"Sent directory to trash: {dirpath}")
            return {
                "success": True, 
                "message": f"Dossier envoye a la corbeille: {dirpath} - RECUPERABLE",
                "dirpath": dirpath,
                "deleted": True,
                "method": "trash"
            }
        else:  # SUPPRESSION DÉFINITIVE
            os.rmdir(dirpath)
            logger.info(f"Permanently deleted directory: {dirpath}")
            return {
                "success": True, 
                "message": f"Dossier supprime definitivement: {dirpath} - NON RECUPERABLE",
                "dirpath": dirpath,
                "deleted": True,
                "method": "permanent"
            }
    except PermissionError:
        return {"success": False, "message": f"Permission denied: {dirpath}", "dirpath": dirpath, "deleted": False}
    except Exception as e:
        logger.error(f"Error deleting directory: {str(e)}")
        return {"success": False, "message": f"Error deleting directory: {str(e)}", "dirpath": dirpath, "deleted": False}


# ---- Resources ----

@server.resource("capabilities://")
def get_capabilities() -> Dict[str, Any]:
    """Provide information about this MCP server's capabilities."""
    return {
        "name": "Document Operations", "version": "0.3.1",
        "description": "MCP server for document operations (Word, Excel, PDF) with file management",
        "document_operations": {
            "word": {"create": True, "create_formatted": True, "edit": True, "convert_from_txt": True, "analyze_structure": True, "compare": True},
            "excel": {"create": True, "edit": True, "convert_from_csv": True},
            "pdf": {"create": True, "convert_from_word": True},
            "file_management": {"delete_file": True, "delete_directory": True, "trash_support": True}
        }
    }


def main():
    """Main entry point for the server."""
    try:
        log_dir = Path(__file__).parent.parent / "logs"
        log_dir.mkdir(exist_ok=True)
        startup_logger = logging.getLogger("startup")
        startup_logger.setLevel(logging.INFO)
        startup_logger.propagate = False
        file_handler = logging.FileHandler(log_dir / "startup.log")
        file_handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
        startup_logger.addHandler(file_handler)
        startup_logger.info("Starting Document Operations MCP Server v0.3.1...")
        server.run()
    except Exception as e:
        logger.error(f"Error starting server: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
